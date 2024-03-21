from docx import Document
from docx.shared import Inches

class WordDriver():
    def __init__(self):
        pass

    def write_content(self, writer, filename):
        document = Document()
        document.add_heading('分子坐标', 0)
        for file_cont in writer:
            for line in file_cont:
                document.add_paragraph(line)
        document.save(filename)

    def write_table(self, writer, filename):
        document = Document()
        document.add_heading('分子频率', 0)

        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = writer[0][0]
        hdr_cells[1].text = writer[0][1]

        for content in writer[1:]:
            row_cells = table.add_row().cells
            row_cells[0].text = content[0]
            row_cells[1].text = "\n".join(content[1:])
        document.save(filename)
 
    def read_content(self, filename):
        doc = Document(filename)
        TextList = []
        for paragraph in doc.paragraphs:
            TextList.append(paragraph.text)
        
        return '\n'.join(TextList)

    def write_single_content(self, writer, filename):
        document = Document()
        document.add_heading('分子坐标', 0)
        for file_cont in writer:
            document.add_paragraph(file_cont)
        document.save(filename)

if __name__ == '__main__':
    A = WordDriver()
    # records = [
    # ["3", '101'],
    # ["7", '422'],
    # ["4", '631']]
    # A.write_content(records, "测试内容.docx")
    # A.write_table(records, "测试表格.docx")
    
    b = A.read_content(r"C:\Users\DELL\Desktop\C4xianan\gas\extract data\gas坐标.docx")
    print(b)
    A.write_single_content(b, "测试内容.docx")